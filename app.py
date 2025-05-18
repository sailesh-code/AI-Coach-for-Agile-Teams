from flask import Flask, jsonify, request, send_file
from flask_cors import CORS
from jira import JIRA
import os
from dotenv import load_dotenv
import google.generativeai as genai
from datetime import datetime, timedelta
import pytz
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import pandas as pd
import numpy as np

# Load environment variables
load_dotenv()

app = Flask(__name__)
CORS(app)

# Configure Gemini
genai.configure(api_key=os.getenv('GEMINI_API_KEY'))
model = genai.GenerativeModel('gemini-2.0-flash')

# Jira configuration
JIRA_URL = os.getenv('JIRA_URL')
JIRA_EMAIL = os.getenv('JIRA_EMAIL')
JIRA_API_TOKEN = os.getenv('JIRA_API_TOKEN')

def get_jira_client():
    if not all([JIRA_URL, JIRA_EMAIL, JIRA_API_TOKEN]):
        raise ValueError("Missing Jira configuration. Please check your .env file.")
    
    # Ensure the URL doesn't end with a slash
    jira_url = JIRA_URL.rstrip('/')
    
    try:
        return JIRA(
            server=jira_url,
            basic_auth=(JIRA_EMAIL, JIRA_API_TOKEN),
            validate=True
        )
    except Exception as e:
        raise Exception(f"Failed to connect to Jira: {str(e)}")

@app.route('/api/boards', methods=['GET'])
def get_boards():
    try:
        jira_client = get_jira_client()
        boards = jira_client.boards()
        return jsonify([{
            'id': board.id,
            'name': board.name,
            'type': board.type
        } for board in boards])
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/sprints', methods=['GET'])
def get_sprints():
    try:
        board_id = request.args.get('boardId')
        if not board_id:
            return jsonify({'error': 'Board ID is required'}), 400

        jira_client = get_jira_client()
        sprints = jira_client.sprints(board_id)
        
        # Format sprint data
        formatted_sprints = []
        for sprint in sprints:
            formatted_sprints.append({
                'id': sprint.id,
                'name': sprint.name,
                'state': sprint.state,
                'startDate': sprint.startDate,
                'endDate': sprint.endDate,
                'goal': sprint.goal if hasattr(sprint, 'goal') else None
            })
        
        # Sort sprints by end date (most recent first)
        formatted_sprints.sort(key=lambda x: x['endDate'] if x['endDate'] else datetime.min, reverse=True)
        
        return jsonify(formatted_sprints)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def get_sprint_stories(jira_client, sprint_id):
    # JQL query to get all stories in the sprint
    jql = f'sprint = {sprint_id} AND type in (Story, Task, Bug) ORDER BY created DESC'
    issues = jira_client.search_issues(jql, maxResults=False, expand='changelog,renderedFields')
    
    stories = []
    for issue in issues:
        # Get all available fields
        story_data = {
            'key': issue.key,
            'summary': issue.fields.summary,
            'description': issue.fields.description,
            'status': issue.fields.status.name,
            'type': issue.fields.issuetype.name,
            'priority': getattr(issue.fields, 'priority', None).name if getattr(issue.fields, 'priority', None) else None,
            'assignee': getattr(issue.fields.assignee, 'displayName', None) if issue.fields.assignee else None,
            'reporter': getattr(issue.fields.reporter, 'displayName', None) if issue.fields.reporter else None,
            'created': issue.fields.created,
            'updated': issue.fields.updated,
            'resolution': getattr(issue.fields.resolution, 'name', None) if issue.fields.resolution else None,
            'labels': getattr(issue.fields, 'labels', []),
            'components': [comp.name for comp in getattr(issue.fields, 'components', [])],
            'story_points': getattr(issue.fields, 'customfield_10016', None),  # Adjust field ID based on your Jira setup
            'epic_link': getattr(issue.fields, 'customfield_10014', None),  # Adjust field ID based on your Jira setup
            'subtasks': [],
            'changelog': [],
            'comments': [],
            'blockers': []
        }
        
        # Get subtasks
        subtasks = jira_client.search_issues(f'parent = {issue.key}', expand='changelog')
        for subtask in subtasks:
            subtask_data = {
                'key': subtask.key,
                'summary': subtask.fields.summary,
                'description': subtask.fields.description,
                'status': subtask.fields.status.name,
                'assignee': getattr(subtask.fields.assignee, 'displayName', None) if subtask.fields.assignee else None,
                'created': subtask.fields.created,
                'updated': subtask.fields.updated,
                'changelog': [],
                'blockers': []
            }
            
            # Get subtask changelog
            for history in subtask.changelog.histories:
                for item in history.items:
                    subtask_data['changelog'].append({
                        'date': history.created,
                        'author': history.author.displayName,
                        'field': item.field,
                        'from': item.fromString,
                        'to': item.toString
                    })
            
            # Get subtask blockers using issue links
            try:
                blocker_links = jira_client.search_issues(f'issue in linkedIssues({subtask.key}) AND type in (Bug, Story, Task)')
                for blocker in blocker_links:
                    subtask_data['blockers'].append({
                        'key': blocker.key,
                        'summary': blocker.fields.summary,
                        'status': blocker.fields.status.name
                    })
            except Exception as e:
                print(f"Error fetching blockers for subtask {subtask.key}: {str(e)}")
            
            story_data['subtasks'].append(subtask_data)
        
        # Get story changelog
        for history in issue.changelog.histories:
            for item in history.items:
                story_data['changelog'].append({
                    'date': history.created,
                    'author': history.author.displayName,
                    'field': item.field,
                    'from': item.fromString,
                    'to': item.toString
                })
        
        # Get comments
        if hasattr(issue.fields, 'comment'):
            for comment in issue.fields.comment.comments:
                story_data['comments'].append({
                    'author': comment.author.displayName,
                    'body': comment.body,
                    'created': comment.created
                })
        
        # Get story blockers using issue links
        try:
            blocker_links = jira_client.search_issues(f'issue in linkedIssues({issue.key}) AND type in (Bug, Story, Task)')
            for blocker in blocker_links:
                story_data['blockers'].append({
                    'key': blocker.key,
                    'summary': blocker.fields.summary,
                    'status': blocker.fields.status.name
                })
        except Exception as e:
            print(f"Error fetching blockers for story {issue.key}: {str(e)}")
        
        stories.append(story_data)
    
    return stories

def generate_subgoals(sprint_goal):
    prompt = f"""
    Do not summarize, rewrite, or rephrase any part of the text. Each subgoal should be exactly as it appears in the original sprint goal, just separated out clearly. Do not make up any subgoals. Do not split any sentences.
    Sprint Goal: {sprint_goal}
    """
    
    response = model.generate_content(prompt)
    return response.text

def assign_stories_to_subgoals(stories, subgoals):
    # Create a detailed prompt for the AI to assign stories to subgoals
    stories_text = "\n".join([
        f"Story {story['key']}:\n"
        f"Summary: {story['summary']}\n"
        f"Description: {story['description']}\n"
        f"Type: {story['type']}\n"
        f"Status: {story['status']}\n"
        f"Labels: {', '.join(story['labels'])}\n"
        f"Components: {', '.join(story['components'])}\n"
        f"Priority: {story['priority']}\n"
        for story in stories
    ])
    
    prompt = f"""
    You are a Product Owner analyzing stories from Jira. Below is the list of user stories and sprint goals for the current sprint.

    Your task is to assign each story to the most relevant sprint goal based on the story's description and summary and acceptance criteria. If a story could relate to multiple goals, assign it to the most appropriate primary goal. If a story does not relate to any of the goals, mark it as "Unassigned".
    
    Subgoals:
    {subgoals}
    
    Stories:
    {stories_text}
    
    Return the assignments in this format:
    Subgoal 1:
    - STORY-123: Story Summary
    - STORY-456: Story Summary

    Subgoal 2:
    - STORY-789: Story Summary
    
    Unassigned:
    - STORY-ABC: Story Summary
    """
    
    response = model.generate_content(prompt)
    return response.text

def generate_achievements(stories, subgoals):
    # Create a detailed prompt for analyzing stories and generating achievements
    stories_text = "\n".join([
        f"Story {story['key']}:\n"
        f"Summary: {story['summary']}\n"
        f"Description: {story['description']}\n"
        f"Type: {story['type']}\n"
        f"Status: {story['status']}\n"
        f"Labels: {', '.join(story['labels'])}\n"
        f"Components: {', '.join(story['components'])}\n"
        f"Priority: {story['priority']}\n"
        f"Comments: {len(story['comments'])} comments\n"
        for story in stories
    ])
    
    prompt = f"""
    Analyze the following stories and their assignments to subgoals. For each subgoal, identify key achievements.
    Focus on completed tasks, technical achievements, improvements, and measurable results.
    
    Subgoals:
    {subgoals}
    
    Stories:
    {stories_text}
    
    For each subgoal, list key achievements as bullet points.
    For each subgoal, list all the assigned stories as story numbers from subgoals variable.
    Format exactly like this:
    Subgoal 1:
    Story Numbers: STORY-123, STORY-456
    - First achievement
    - Second achievement
    - Third achievement
    
    Subgoal 2:
    Story Numbers: STORY-789
    - First achievement
    - Second achievement
    - Third achievement
    
    Important: Each achievement should be a complete sentence starting with a capital letter.
    Do not number the achievements or add any additional formatting.
    """
    
    response = model.generate_content(prompt)
    return response.text

@app.route('/api/sprint-report', methods=['GET'])
def get_sprint_report():
    try:
        board_id = request.args.get('boardId')
        sprint_id = request.args.get('sprintId')
        
        if not board_id or not sprint_id:
            return jsonify({'error': 'Board ID and Sprint ID are required'}), 400

        jira_client = get_jira_client()
        
        # Get sprint details
        sprint = jira_client.sprint(sprint_id)
        if not sprint:
            return jsonify({'error': 'Sprint not found'}), 404
        
        # Get sprint goal
        sprint_goal = sprint.goal if hasattr(sprint, 'goal') else "No sprint goal found"
        
        # Generate subgoals using Gemini
        subgoals = generate_subgoals(sprint_goal)
        
        # Get stories from the sprint
        stories = get_sprint_stories(jira_client, sprint_id)
        
        # Assign stories to subgoals
        story_assignments = assign_stories_to_subgoals(stories, subgoals)
        
        # Generate achievements for each subgoal
        achievements = generate_achievements(stories, subgoals)
        
        return jsonify({
            'sprint_name': sprint.name,
            'sprint_goal': sprint_goal,
            'subgoals': subgoals,
            'stories': stories,
            'story_assignments': story_assignments,
            'achievements': achievements,
            'start_date': sprint.startDate,
            'end_date': sprint.endDate
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/sprint-report/download', methods=['GET'])
def download_sprint_report():
    try:
        board_id = request.args.get('boardId')
        sprint_id = request.args.get('sprintId')
        
        if not board_id or not sprint_id:
            return jsonify({'error': 'Board ID and Sprint ID are required'}), 400

        jira_client = get_jira_client()
        
        # Get sprint details
        sprint = jira_client.sprint(sprint_id)
        if not sprint:
            return jsonify({'error': 'Sprint not found'}), 404
        
        # Get sprint goal
        sprint_goal = sprint.goal if hasattr(sprint, 'goal') else "No sprint goal found"
        
        # Generate subgoals using Gemini
        subgoals = generate_subgoals(sprint_goal)
        
        # Get stories from the sprint
        stories = get_sprint_stories(jira_client, sprint_id)
        
        # Assign stories to subgoals
        story_assignments = assign_stories_to_subgoals(stories, subgoals)
        
        # Generate achievements for each subgoal
        achievements = generate_achievements(stories, subgoals)

        # Create a new Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Sprint Report: {sprint.name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add sprint dates
        dates = doc.add_paragraph()
        dates.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dates.add_run(f'{sprint.startDate} - {sprint.endDate}').italic = True
        
        # Add sprint goal
        doc.add_heading('Sprint Goal', level=1)
        doc.add_paragraph(sprint_goal)
        
        # Add achievements and story assignments
        doc.add_heading('Achievements and Story Assignments', level=1)
        
        # Process achievements and story assignments
        achievement_sections = achievements.split('\n\n')
        assignment_sections = story_assignments.split('\n\n')
        
        for section in achievement_sections:
            if not section.strip():
                continue
                
            lines = section.split('\n')
            if not lines:
                continue
                
            # Add subgoal heading
            subgoal_heading = doc.add_heading(lines[0], level=2)
            
            # Add story numbers if available
            story_numbers = next((line for line in lines if line.startswith('Story Numbers:')), None)
            if story_numbers:
                doc.add_paragraph(story_numbers)
            
            # Add achievements subheading
            doc.add_heading('Achievements', level=3)
            
            # Add achievements
            for line in lines:
                if line.startswith('- '):
                    p = doc.add_paragraph()
                    p.style = 'List Bullet'
                    p.add_run(line[2:])
            
            # Add a small space between subgoals
            doc.add_paragraph()
        
        # Save the document to a BytesIO object
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return send_file(
            doc_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'sprint_report_{sprint.name}.docx'
        )
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def process_excel_data(excel_file):
    """Process Excel file and extract relevant information using LLM."""
    try:
        # Read Excel file
        df = pd.read_excel(excel_file)
        
        # Convert DataFrame to string representation
        excel_data = df.to_string()
        
        # Create prompt for LLM to extract structured data
        prompt = f"""
        Analyze the following Excel data and extract the following information in a structured format:
        1. Sprint Capacity
        2. Team Member Capacities
        3. Story Details (including subtasks)
        4. Changelogs
        5. Blockers/Impediments
        
        Excel Data:
        {excel_data}
        
        Return the data in this JSON format:
        {{
            "sprint_capacity": {{
                "total_capacity": number,
                "unit": "hours/days"
            }},
            "team_members": [
                {{
                    "name": string,
                    "capacity": number,
                    "unit": "hours/days"
                }}
            ],
            "stories": [
                {{
                    "id": string,
                    "summary": string,
                    "description": string,
                    "status": string,
                    "assignee": string,
                    "subtasks": [
                        {{
                            "id": string,
                            "summary": string,
                            "status": string,
                            "assignee": string
                        }}
                    ],
                    "changelog": [
                        {{
                            "date": string,
                            "field": string,
                            "from": string,
                            "to": string
                        }}
                    ],
                    "blockers": [
                        {{
                            "description": string,
                            "resolution": string
                        }}
                    ]
                }}
            ]
        }}
        
        Important: Return ONLY the JSON object, with no additional text or explanation.
        """
        
        response = model.generate_content(prompt)
        
        # Clean the response text to ensure it's valid JSON
        response_text = response.text.strip()
        
        # Remove any markdown code block indicators if present
        if response_text.startswith('```json'):
            response_text = response_text[7:]
        if response_text.startswith('```'):
            response_text = response_text[3:]
        if response_text.endswith('```'):
            response_text = response_text[:-3]
            
        response_text = response_text.strip()
        
        try:
            structured_data = json.loads(response_text)
            return structured_data
        except json.JSONDecodeError as e:
            # If JSON parsing fails, try to extract JSON from the response
            import re
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if json_match:
                try:
                    structured_data = json.loads(json_match.group())
                    return structured_data
                except json.JSONDecodeError:
                    raise Exception(f"Failed to parse JSON from response: {str(e)}")
            else:
                raise Exception(f"Failed to extract JSON from response: {str(e)}")
                
    except Exception as e:
        raise Exception(f"Error processing Excel file: {str(e)}")

def parse_jira_datetime(datetime_str):
    """Parse Jira datetime string to datetime object."""
    try:
        if not datetime_str:
            print("Empty datetime string received")
            return None
            
        print(f"Parsing datetime: {datetime_str}")
        
        # Handle different Jira datetime formats
        if '+' in datetime_str:
            # Format: 2025-05-16T15:38:57.738+0530
            dt_str, offset = datetime_str.split('+')
            dt = datetime.fromisoformat(dt_str)
            hours = int(offset[:2])
            minutes = int(offset[2:4])
            dt = dt - timedelta(hours=hours, minutes=minutes)
        elif 'Z' in datetime_str:
            # Format: 2025-05-16T15:38:57.738Z
            dt = datetime.fromisoformat(datetime_str.replace('Z', '+00:00'))
        else:
            # Format: 2025-05-16T15:38:57.738
            dt = datetime.fromisoformat(datetime_str)
            
        print(f"Successfully parsed datetime: {dt}")
        return dt
    except Exception as e:
        print(f"Error parsing datetime {datetime_str}: {str(e)}")
        return None

def analyze_sprint_churn(sprint_data):
    """Analyze sprint churn by examining story changes during the sprint."""
    print(f"Analyzing sprint data: {json.dumps(sprint_data, indent=2)}")
    
    # Extract and validate sprint dates
    sprint_start_str = sprint_data.get('start_date')
    sprint_end_str = sprint_data.get('end_date')
    
    if not sprint_start_str or not sprint_end_str:
        print(f"Missing sprint dates. Start: {sprint_start_str}, End: {sprint_end_str}")
        raise Exception("Missing sprint dates")
    
    sprint_start = parse_jira_datetime(sprint_start_str)
    sprint_end = parse_jira_datetime(sprint_end_str)
    
    if not sprint_start:
        print(f"Failed to parse sprint start date: {sprint_start_str}")
        raise Exception(f"Invalid sprint start date: {sprint_start_str}")
    if not sprint_end:
        print(f"Failed to parse sprint end date: {sprint_end_str}")
        raise Exception(f"Invalid sprint end date: {sprint_end_str}")
    
    print(f"Sprint period: {sprint_start} to {sprint_end}")
    
    churn_analysis = {
        'added_stories': [],
        'removed_stories': [],
        'modified_stories': [],
        'status_changes': {},
        'story_point_changes': {},
        'assignee_changes': {}
    }
    
    for story in sprint_data['stories']:
        story_changes = {
            'key': story['key'],
            'summary': story['summary'],
            'changes': [],
            'status_changes': [],
            'point_changes': [],
            'assignee_changes': []
        }
        
        # Analyze changelog entries within sprint dates
        for change in story['changelog']:
            change_date = parse_jira_datetime(change['date'])
            if change_date and sprint_start <= change_date <= sprint_end:
                if change['field'] == 'status':
                    story_changes['status_changes'].append({
                        'date': change['date'],
                        'from': change['from'],
                        'to': change['to']
                    })
                elif change['field'] == 'Story Points':
                    story_changes['point_changes'].append({
                        'date': change['date'],
                        'from': change['from'],
                        'to': change['to']
                    })
                elif change['field'] == 'assignee':
                    story_changes['assignee_changes'].append({
                        'date': change['date'],
                        'from': change['from'],
                        'to': change['to']
                    })
                
                story_changes['changes'].append({
                    'date': change['date'],
                    'field': change['field'],
                    'from': change['from'],
                    'to': change['to']
                })
        
        # Check if story was added during sprint
        story_created = parse_jira_datetime(story['created'])
        if story_created and sprint_start <= story_created <= sprint_end:
            churn_analysis['added_stories'].append({
                'key': story['key'],
                'summary': story['summary'],
                'created_date': story['created']
            })
        
        # If story has changes during sprint, add to modified stories
        if story_changes['changes']:
            churn_analysis['modified_stories'].append(story_changes)
            
            # Track status changes
            if story_changes['status_changes']:
                churn_analysis['status_changes'][story['key']] = story_changes['status_changes']
            
            # Track story point changes
            if story_changes['point_changes']:
                churn_analysis['story_point_changes'][story['key']] = story_changes['point_changes']
            
            # Track assignee changes
            if story_changes['assignee_changes']:
                churn_analysis['assignee_changes'][story['key']] = story_changes['assignee_changes']
    
    return churn_analysis

def analyze_churned_stories(sprint_data):
    """Analyze stories that were added to the sprint after it started."""
    churned_stories = []
    sprint_start = parse_jira_datetime(sprint_data['start_date'])
    sprint_end = parse_jira_datetime(sprint_data['end_date'])
    
    if not sprint_start or not sprint_end:
        raise Exception("Invalid sprint dates")
    
    # Ensure both sprint dates are timezone-aware
    utc = pytz.UTC
    if sprint_start.tzinfo is None:
        sprint_start = utc.localize(sprint_start)
    if sprint_end.tzinfo is None:
        sprint_end = utc.localize(sprint_end)
    
    for story in sprint_data['stories']:
        # Check changelog for sprint changes
        for change in story['changelog']:
            if change['field'] == 'Sprint':
                change_date = parse_jira_datetime(change['date'])
                if change_date:
                    # Ensure change date is timezone-aware
                    if change_date.tzinfo is None:
                        change_date = utc.localize(change_date)
                    
                    if sprint_start <= change_date <= sprint_end:
                        # Story was added to this sprint during the sprint
                        churned_stories.append({
                            'story_id': story['key'],
                            'summary': story['summary'],
                            'added_date': change['date'],
                            'status': story['status'],
                            'assignee': story['assignee'],
                            'story_points': story['story_points'],
                            'type': story['type']
                        })
    
    # Calculate churn metrics
    total_churned = len(churned_stories)
    total_churned_points = sum(story['story_points'] for story in churned_stories if story['story_points'])
    
    churn_by_type = {
        'Story': {
            'count': len([s for s in churned_stories if s['type'] == 'Story']),
            'points': sum(s['story_points'] for s in churned_stories if s['type'] == 'Story' and s['story_points'])
        },
        'Task': {
            'count': len([s for s in churned_stories if s['type'] == 'Task']),
            'points': sum(s['story_points'] for s in churned_stories if s['type'] == 'Task' and s['story_points'])
        },
        'Bug': {
            'count': len([s for s in churned_stories if s['type'] == 'Bug']),
            'points': sum(s['story_points'] for s in churned_stories if s['type'] == 'Bug' and s['story_points'])
        }
    }
    
    return {
        'churned_stories': churned_stories,
        'total_churned': total_churned,
        'total_churned_points': total_churned_points,
        'churn_by_type': churn_by_type,
        'churn_summary': f"Total Churned Stories: {total_churned} ({total_churned_points} points)\n" +
                        f"Stories: {churn_by_type['Story']['count']} ({churn_by_type['Story']['points']} points)\n" +
                        f"Tasks: {churn_by_type['Task']['count']} ({churn_by_type['Task']['points']} points)\n" +
                        f"Bugs: {churn_by_type['Bug']['count']} ({churn_by_type['Bug']['points']} points)"
    }

def analyze_subgoal_improvements(stories, subgoal):
    """Analyze improvement areas for a specific subgoal based on its stories."""
    # Create a detailed prompt for analyzing stories under a subgoal
    stories_text = "\n".join([
        f"Story {story['key']}:\n"
        f"Summary: {story['summary']}\n"
        f"Description: {story['description']}\n"
        f"Status: {story['status']}\n"
        f"Comments: {len(story['comments'])} comments\n"
        f"Changelog Entries: {len(story['changelog'])} entries\n"
        f"Blockers: {len(story['blockers'])} blockers\n"
        for story in stories
    ])
    
    prompt = f"""
    Analyze the following stories under this subgoal and identify specific improvement areas. Focus on concrete, actionable improvements based on:
    1. Story Summary and Description
       - Clarity and completeness
       - Acceptance criteria
       - Technical requirements
    2. Status Changes
       - Time spent in each status
       - Status transition patterns
    3. Comments
       - Communication effectiveness
       - Knowledge sharing
       - Decision documentation
    4. Changelog Entries
       - Frequency of changes
       - Types of changes
       - Impact on delivery
    5. Blockers
       - Nature of blockers
       - Resolution time
       - Prevention opportunities

    Subgoal: {subgoal}
    
    Stories:
    {stories_text}
    
    For each improvement area, provide:
    1. Specific observation from the data
    2. Concrete impact on the sprint
    3. Actionable recommendation
    
    Return the analysis in this format:
    {{
        "improvement_areas": [
            {{
                "category": "Story Definition/Status/Comments/Changelog/Blockers",
                "observation": "Specific observation from the data",
                "impact": "Concrete impact on the sprint",
                "recommendation": "Actionable recommendation"
            }}
        ]
    }}
    
    Important: 
    - Be specific and data-driven
    - Focus on concrete improvements, not general suggestions
    - Base recommendations on actual patterns in the data
    - Return ONLY the JSON object
    """
    
    response = model.generate_content(prompt)
    
    # Clean the response text to ensure it's valid JSON
    response_text = response.text.strip()
    
    # Remove any markdown code block indicators if present
    if response_text.startswith('```json'):
        response_text = response_text[7:]
    if response_text.startswith('```'):
        response_text = response_text[3:]
    if response_text.endswith('```'):
        response_text = response_text[:-3]
        
    response_text = response_text.strip()
    
    try:
        return json.loads(response_text)
    except json.JSONDecodeError as e:
        # If JSON parsing fails, try to extract JSON from the response
        import re
        json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
        if json_match:
            try:
                return json.loads(json_match.group())
            except json.JSONDecodeError:
                raise Exception(f"Failed to parse JSON from response: {str(e)}")
        else:
            raise Exception(f"Failed to extract JSON from response: {str(e)}")

def calculate_spillover_points(sprint_data, spilled_stories):
    """Calculate story points for stories that spilled over from the sprint."""
    # Get sprint dates
    sprint_start = parse_jira_datetime(sprint_data['start_date'])
    sprint_end = parse_jira_datetime(sprint_data['end_date'])
    
    # Ensure both sprint dates are timezone-aware
    utc = pytz.UTC
    if sprint_start and sprint_start.tzinfo is None:
        sprint_start = utc.localize(sprint_start)
    if sprint_end and sprint_end.tzinfo is None:
        sprint_end = utc.localize(sprint_end)
    
    total_spilled_points = 0
    spilled_stories_with_points = []
    
    for story in spilled_stories:
        story_id = story['story_id']
        # Find the corresponding story in sprint_data
        original_story = next((s for s in sprint_data['stories'] if s['key'] == story_id), None)
        if not original_story:
            continue
            
        # Check if story was in sprint at start
        was_in_sprint_at_start = False
        for change in original_story['changelog']:
            change_date = parse_jira_datetime(change['date'])
            if not change_date:
                continue
                
            # Ensure change_date is timezone-aware
            if change_date.tzinfo is None:
                change_date = utc.localize(change_date)
            
            if change['field'] == 'Sprint' and change_date <= sprint_start:
                was_in_sprint_at_start = True
                break
        
        # If no sprint changes found, check creation date
        if not was_in_sprint_at_start and not any(change['field'] == 'Sprint' for change in original_story['changelog']):
            story_created = parse_jira_datetime(original_story['created'])
            if story_created:
                if story_created.tzinfo is None:
                    story_created = utc.localize(story_created)
                if story_created <= sprint_start:
                    was_in_sprint_at_start = True
        
        # Only count points if story was in sprint at start
        if was_in_sprint_at_start:
            story_points = original_story.get('story_points', 0) or 0  # Convert None to 0
            total_spilled_points += story_points
            spilled_stories_with_points.append({
                'story_id': story_id,
                'story_points': story_points,
                'reason': story['reason']
            })
    
    return total_spilled_points, spilled_stories_with_points

def generate_improvement_areas(structured_data, sprint_data):
    """Generate improvement areas using LLM based on structured data and sprint data."""
    # First analyze churned stories
    churn_analysis = analyze_churned_stories(sprint_data)
    
    prompt = f"""
    Analyze the following sprint data and generate detailed improvement areas. Focus on:
    1. Spill-over Analysis
       - Identify stories that spilled over
       - Identify stories from changelog if their sprint number has changed from current sprint to future sprints or removed from current sprint during the sprint.
       - Analyze root causes
       - Suggest preventive measures
    
    2. Churn Analysis
       - Analyze the following churned stories that were added to the sprint after it started:
       {json.dumps(churn_analysis, indent=2)}
       - Identify all churned stories that were added to the sprint after it started.
       - Analyze impact on sprint velocity considering both story count and story points
       - Suggest ways to reduce churn based on the type and size of churned stories
    
    3. Team Utilization
       - Calculate utilization for each team member:
         * Get their story point capacity for the sprint
         * Count total story points completed by them, only consider the story points completed during the sprint. Do not consider the story points completed before the sprint start date and after the sprint end date. The story must be marked as done during the sprint.
         * Calculate utilization as: (completed story points / capacity) * 100
       - Consider a team member over-utilized if:
         * They complete more story points than their capacity, only consider the story points completed during the sprint. Do not consider the story points completed before the sprint start date and after the sprint end date. The story must be marked as done during the sprint.
       - Consider a team member under-utilized if:
         * They handle less story points than their capacity, only consider the story points completed during the sprint. Do not consider the story points completed before the sprint start date and after the sprint end date. The story must be marked as done during the sprint.
       - Analyze workload distribution
       - Suggest optimal resource allocation
    
    4. Additional Improvement Areas
       - Identify any other patterns or issues
       - Suggest specific actionable improvements
    
    Sprint Data:
    {json.dumps(sprint_data, indent=2)}
    
    Excel Data:
    {json.dumps(structured_data, indent=2)}
    
    Return the analysis in this JSON format:
    {{
        "spill_over_analysis": {{
            "spilled_stories": [
                {{
                    "story_id": string,
                    "reason": string,
                    "prevention_suggestion": string
                }}
            ],
            "root_causes": [string],
            "recommendations": [string]
        }},
        "churn_analysis": {{
            "high_churn_stories": [
                {{
                    "story_id": string,
                    "churn_count": number,
                    "story_points": number,
                    "impact": string
                }}
            ],
            "velocity_impact": string,
            "reduction_suggestions": [string]
        }},
        "team_utilization": {{
            "under_utilized": [
                {{
                    "member": string,
                    "capacity": number,
                    "completed_points": number,
                    "utilization": number,
                    "suggestion": string
                }}
            ],
            "over_utilized": [
                {{
                    "member": string,
                    "capacity": number,
                    "completed_points": number,
                    "utilization": number,
                    "suggestion": string
                }}
            ],
            "workload_distribution": string,
            "optimization_suggestions": [string]
        }},
        "additional_improvements": [
            {{
                "area": string,
                "observation": string,
                "suggestion": string
            }}
        ]
    }}
    
    Important: Return ONLY the JSON object, with no additional text or explanation.
    """
    
    response = model.generate_content(prompt)
    
    # Clean the response text to ensure it's valid JSON
    response_text = response.text.strip()
    
    # Remove any markdown code block indicators if present
    if response_text.startswith('```json'):
        response_text = response_text[7:]
    if response_text.startswith('```'):
        response_text = response_text[3:]
    if response_text.endswith('```'):
        response_text = response_text[:-3]
        
    response_text = response_text.strip()
    
    try:
        return json.loads(response_text)
    except json.JSONDecodeError as e:
        # If JSON parsing fails, try to extract JSON from the response
        import re
        json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
        if json_match:
            try:
                return json.loads(json_match.group())
            except json.JSONDecodeError:
                raise Exception(f"Failed to parse JSON from response: {str(e)}")
        else:
            raise Exception(f"Failed to extract JSON from response: {str(e)}")

def calculate_sprint_metrics(sprint_data):
    """Calculate sprint metrics including unassigned stories."""
    # Get sprint dates
    sprint_start = parse_jira_datetime(sprint_data['start_date'])
    sprint_end = parse_jira_datetime(sprint_data['end_date'])
    
    # Ensure both sprint dates are timezone-aware
    utc = pytz.UTC
    if sprint_start and sprint_start.tzinfo is None:
        sprint_start = utc.localize(sprint_start)
    if sprint_end and sprint_end.tzinfo is None:
        sprint_end = utc.localize(sprint_end)
    
    # Initialize metrics
    metrics = {
        'committed': 0,
        'completed': 0
    }
    
    for story in sprint_data['stories']:
        story_points = story['story_points'] or 0
        
        # Track story status
        was_in_sprint_at_start = False
        was_completed_during_sprint = False
        
        # Check changelog to determine story status
        for change in story['changelog']:
            change_date = parse_jira_datetime(change['date'])
            if not change_date:
                continue
                
            # Ensure change_date is timezone-aware
            if change_date.tzinfo is None:
                change_date = utc.localize(change_date)
            
            # Check if story was in sprint at start
            if change['field'] == 'Sprint':
                if change_date <= sprint_start:
                    was_in_sprint_at_start = True
            
            # Check if story was completed during sprint
            if change['field'] == 'status' and change['to'] == 'Done':
                if sprint_start <= change_date <= sprint_end:
                    was_completed_during_sprint = True
        
        # If no sprint changes found, check creation date
        if not any(change['field'] == 'Sprint' for change in story['changelog']):
            story_created = parse_jira_datetime(story['created'])
            if story_created:
                if story_created.tzinfo is None:
                    story_created = utc.localize(story_created)
                if story_created <= sprint_start:
                    was_in_sprint_at_start = True
        
        # Add to committed points if story was in sprint at start
        if was_in_sprint_at_start:
            metrics['committed'] += story_points
        
        # Add to completed points if story was completed during sprint
        if was_completed_during_sprint:
            metrics['completed'] += story_points
    
    return metrics

def calculate_member_story_points(sprint_data):
    """Calculate committed and completed story points for each team member."""
    # Initialize member data
    member_data = {}
    
    # Get sprint dates
    sprint_start = parse_jira_datetime(sprint_data['start_date'])
    sprint_end = parse_jira_datetime(sprint_data['end_date'])
    
    # Ensure both sprint dates are timezone-aware
    utc = pytz.UTC
    if sprint_start and sprint_start.tzinfo is None:
        sprint_start = utc.localize(sprint_start)
    if sprint_end and sprint_end.tzinfo is None:
        sprint_end = utc.localize(sprint_end)
    
    for story in sprint_data['stories']:
        assignee = story['assignee']
        if not assignee:
            continue
            
        # Initialize member data if not exists
        if assignee not in member_data:
            member_data[assignee] = {
                'committed': 0,
                'completed': 0
            }
            
        story_points = story['story_points'] or 0
        
        # Track story status
        was_in_sprint_at_start = False
        was_completed_during_sprint = False
        
        # Check changelog to determine story status
        for change in story['changelog']:
            change_date = parse_jira_datetime(change['date'])
            if not change_date:
                continue
                
            # Ensure change_date is timezone-aware
            if change_date.tzinfo is None:
                change_date = utc.localize(change_date)
            
            # Check if story was in sprint at start
            if change['field'] == 'Sprint':
                if change_date <= sprint_start:
                    was_in_sprint_at_start = True
            
            # Check if story was completed during sprint
            if change['field'] == 'status' and change['to'] == 'Done':
                if sprint_start <= change_date <= sprint_end:
                    was_completed_during_sprint = True
        
        # If no sprint changes found, check creation date
        if not any(change['field'] == 'Sprint' for change in story['changelog']):
            story_created = parse_jira_datetime(story['created'])
            if story_created:
                if story_created.tzinfo is None:
                    story_created = utc.localize(story_created)
                if story_created <= sprint_start:
                    was_in_sprint_at_start = True
        
        # Add to committed points if story was in sprint at start
        if was_in_sprint_at_start:
            member_data[assignee]['committed'] += story_points
        
        # Add to completed points if story was completed during sprint
        if was_completed_during_sprint:
            member_data[assignee]['completed'] += story_points
    
    return member_data

def generate_member_capacity_table(structured_data, sprint_data):
    """Generate a table showing member-wise capacity and utilization using LLM."""
    try:
        print("Starting member capacity table generation...")
        
        # Calculate story points for each member
        member_points = calculate_member_story_points(sprint_data)
        
        # Prepare team capacity data
        team_capacity = "\n".join([
            f"Team Member: {member['name']}, Capacity: {member['capacity']} points"
            for member in structured_data['team_members']
        ])
        
        # Prepare member points data
        member_points_text = "\n".join([
            f"Member: {member}, Committed Points: {data['committed']}, Completed Points: {data['completed']}"
            for member, data in member_points.items()
        ])

        # Create prompt for LLM
        prompt = f"""
        Generate a member-wise capacity table using the following data.
        Consider the sprint period from {sprint_data['start_date']} to {sprint_data['end_date']}.

        Team Capacity:
        {team_capacity}

        Calculated Story Points:
        {member_points_text}

        For each team member, calculate:
        1. Use the provided committed and completed points
        2. Calculate utilization: (Completed points / Capacity) * 100

        Return the data in this JSON format:
        {{
            "members": [
                {{
                    "assignee": "string",
                    "capacity": number,
                    "committed": number,
                    "completed": number,
                    "utilization": "string (percentage with % symbol)"
                }}
            ]
        }}

        Important:
        - Use the exact committed and completed points provided
        - Calculate utilization based on the provided capacity
        - Return ONLY the JSON object, with no additional text or explanation
        """

        print("Sending prompt to LLM...")
        # Get response from LLM
        response = model.generate_content(prompt)
        print("Received response from LLM")
        
        # Clean the response text to ensure it's valid JSON
        response_text = response.text.strip()
        print(f"Raw LLM response: {response_text[:200]}...")  # Print first 200 chars
        
        # Remove any markdown code block indicators if present
        if response_text.startswith('```json'):
            response_text = response_text[7:]
        if response_text.startswith('```'):
            response_text = response_text[3:]
        if response_text.endswith('```'):
            response_text = response_text[:-3]
            
        response_text = response_text.strip()
        print(f"Cleaned response text: {response_text[:200]}...")  # Print first 200 chars
        
        try:
            # Parse the JSON response
            result = json.loads(response_text)
            print("Successfully parsed JSON response")
            return result['members']
        except json.JSONDecodeError as e:
            print(f"JSON parsing error: {str(e)}")
            # If JSON parsing fails, try to extract JSON from the response
            import re
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if json_match:
                try:
                    result = json.loads(json_match.group())
                    print("Successfully extracted and parsed JSON from response")
                    return result['members']
                except json.JSONDecodeError as e2:
                    print(f"Failed to parse extracted JSON: {str(e2)}")
                    raise Exception(f"Failed to parse JSON from LLM response: {str(e2)}")
            else:
                print("No JSON object found in response")
                raise Exception(f"Failed to extract JSON from LLM response: {str(e)}")
    
    except Exception as e:
        import traceback
        print(f"Error in generate_member_capacity_table: {str(e)}")
        print("Full traceback:")
        print(traceback.format_exc())
        raise

def generate_combined_sprint_doc(sprint_data, improvement_areas, subgoals, story_assignments, achievements, structured_data):
    """Generate a Word document containing both sprint report and analysis."""
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'Sprint Report & Analysis: {sprint_data["sprint_name"]}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add sprint details
    doc.add_heading('Sprint Details', level=1)
    doc.add_paragraph(f'Sprint Goal: {sprint_data["sprint_goal"] or "No sprint goal defined"}')
    doc.add_paragraph(f'Start Date: {sprint_data["start_date"]}')
    doc.add_paragraph(f'End Date: {sprint_data["end_date"]}')
    
    # Add Sprint Summary
    doc.add_heading('Sprint Summary', level=1)
    
    # Calculate sprint metrics
    # Sprint Capacity
    total_capacity = sum(member.get('capacity', 0) or 0 for member in structured_data['team_members'])
    
    # Calculate total committed and completed points (including unassigned)
    sprint_metrics = calculate_sprint_metrics(sprint_data)
    total_committed = sprint_metrics['committed']
    total_completed = sprint_metrics['completed']
    
    # Churn
    churned_stories = improvement_areas['churn_analysis']['high_churn_stories']
    total_churned = len(churned_stories)
    total_churned_points = sum(story.get('story_points', 0) or 0 for story in churned_stories)
    
    # Spillover
    spilled_stories = improvement_areas['spill_over_analysis']['spilled_stories']
    total_spilled = len(spilled_stories)
    total_spilled_points, spilled_stories_with_points = calculate_spillover_points(sprint_data, spilled_stories)
    
    # Add metrics table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'Value'
    
    # Add metrics
    metrics = [
        ('Sprint Capacity', f'{total_capacity} points'),
        ('Committed Story Points', f'{total_committed} points'),
        ('Velocity', f'{total_completed} points'),
        ('Churn', f'{total_churned} stories ({total_churned_points} points)'),
        ('Spillover', f'{total_spilled} stories ({total_spilled_points} points)')
    ]
    
    for metric, value in metrics:
        row = table.add_row()
        row.cells[0].text = metric
        row.cells[1].text = value
    
    # Add Member Capacity Table
    doc.add_heading('Team Member Capacity Analysis', level=1)
    member_data = generate_member_capacity_table(structured_data, sprint_data)
    
    # Create member capacity table
    member_table = doc.add_table(rows=1, cols=5)
    member_table.style = 'Table Grid'
    
    # Add headers
    header_cells = member_table.rows[0].cells
    header_cells[0].text = 'Assignee'
    header_cells[1].text = 'Capacity (Points)'
    header_cells[2].text = 'Committed (Points)'
    header_cells[3].text = 'Completed (Points)'
    header_cells[4].text = 'Utilization'
    
    # Add member data
    for member in member_data:
        row = member_table.add_row()
        row.cells[0].text = member['assignee']
        row.cells[1].text = str(member.get('capacity', 0) or 0)
        row.cells[2].text = str(member.get('committed', 0) or 0)
        row.cells[3].text = str(member.get('completed', 0) or 0)
        row.cells[4].text = member.get('utilization', '0%')
    
    # Add Sprint Report Section
    doc.add_heading('Sprint Report', level=1)
    
    # Add subgoals and achievements
    doc.add_heading('Sprint Goals and Achievements', level=2)
    achievement_sections = achievements.split('\n\n')
    for section in achievement_sections:
        if not section.strip():
            continue
            
        lines = section.split('\n')
        if not lines:
            continue
            
        # Add subgoal heading
        subgoal_heading = doc.add_heading(lines[0], level=3)
        
        # Add story numbers if available
        story_numbers = next((line for line in lines if line.startswith('Story Numbers:')), None)
        if story_numbers:
            doc.add_paragraph(story_numbers)
            
            # Extract story IDs from the story numbers line
            story_ids = [s.strip() for s in story_numbers.replace('Story Numbers:', '').split(',')]
            
            # Get stories for this subgoal
            subgoal_stories = [s for s in sprint_data['stories'] if s['key'] in story_ids]
            
            # Analyze improvements for this subgoal
            improvements = analyze_subgoal_improvements(subgoal_stories, lines[0])
        
        # Add achievements subheading
        doc.add_heading('Achievements', level=4)
        
        # Add achievements
        for line in lines:
            if line.startswith('- '):
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                p.add_run(line[2:])
        
        # Add improvement areas if available
        if story_numbers and improvements.get('improvement_areas'):
            doc.add_heading('Improvement Areas', level=4)
            
            # Group improvements by category
            improvements_by_category = {}
            for imp in improvements['improvement_areas']:
                category = imp['category']
                if category not in improvements_by_category:
                    improvements_by_category[category] = []
                improvements_by_category[category].append(imp)
            
            # Add improvements by category
            for category, category_improvements in improvements_by_category.items():
                doc.add_heading(category, level=5)
                for imp in category_improvements:
                    p = doc.add_paragraph()
                    p.add_run('Observation: ').bold = True
                    p.add_run(imp['observation'])
                    p = doc.add_paragraph()
                    p.add_run('Impact: ').bold = True
                    p.add_run(imp['impact'])
                    p = doc.add_paragraph()
                    p.add_run('Recommendation: ').bold = True
                    p.add_run(imp['recommendation'])
                    doc.add_paragraph()  # Add spacing between improvements
        
        # Add a small space between subgoals
        doc.add_paragraph()
    
    # Add Sprint Analysis Section
    doc.add_heading('Sprint Analysis', level=1)
    
    # Add Spill-over Analysis
    doc.add_heading('Spill-over Analysis', level=2)
    if spilled_stories_with_points:
        doc.add_paragraph('Spilled Stories:')
        for story in spilled_stories_with_points:
            p = doc.add_paragraph()
            p.add_run(f'Story ID: {story["story_id"]}\n').bold = True
            p.add_run(f'Story Points: {story["story_points"]}\n')
            p.add_run(f'Reason: {story["reason"]}')
    else:
        doc.add_paragraph('No stories spilled over in this sprint.')
    
    doc.add_paragraph('Root Causes:')
    for cause in improvement_areas['spill_over_analysis']['root_causes']:
        doc.add_paragraph(cause, style='List Bullet')
    
    doc.add_paragraph('Recommendations:')
    for rec in improvement_areas['spill_over_analysis']['recommendations']:
        doc.add_paragraph(rec, style='List Bullet')
    
    # Add Churn Analysis
    doc.add_heading('Churn Analysis', level=2)
    if improvement_areas['churn_analysis']['high_churn_stories']:
        # Add churn summary table
        churn_table = doc.add_table(rows=1, cols=3)
        churn_table.style = 'Table Grid'
        
        # Add headers
        header_cells = churn_table.rows[0].cells
        header_cells[0].text = 'Story ID'
        header_cells[1].text = 'Story Points'
        header_cells[2].text = 'Impact'
        
        # Add churned stories
        for story in improvement_areas['churn_analysis']['high_churn_stories']:
            row = churn_table.add_row()
            row.cells[0].text = story['story_id']
            row.cells[1].text = str(story.get('story_points', 0) or 0)
            row.cells[2].text = story.get('impact', '')
    else:
        doc.add_paragraph('No high churn stories identified in this sprint.')
    
    doc.add_paragraph(f'Velocity Impact: {improvement_areas["churn_analysis"]["velocity_impact"]}')
    
    doc.add_paragraph('Reduction Suggestions:')
    for suggestion in improvement_areas['churn_analysis']['reduction_suggestions']:
        doc.add_paragraph(suggestion, style='List Bullet')
    
    # Add Team Utilization
    doc.add_heading('Team Utilization', level=2)
    
    # Add utilization summary table
    util_table = doc.add_table(rows=1, cols=4)
    util_table.style = 'Table Grid'
    
    # Add headers
    header_cells = util_table.rows[0].cells
    header_cells[0].text = 'Team Member'
    header_cells[1].text = 'Capacity'
    header_cells[2].text = 'Completed Points'
    header_cells[3].text = 'Utilization'
    
    # Add over-utilized members
    for member in improvement_areas['team_utilization']['over_utilized']:
        row = util_table.add_row()
        row.cells[0].text = member['member']
        row.cells[1].text = str(member.get('capacity', 0) or 0)
        row.cells[2].text = str(member.get('completed_points', 0) or 0)
        row.cells[3].text = f"{member.get('utilization', 0)}%"
    
    # Add under-utilized members
    for member in improvement_areas['team_utilization']['under_utilized']:
        row = util_table.add_row()
        row.cells[0].text = member['member']
        row.cells[1].text = str(member.get('capacity', 0) or 0)
        row.cells[2].text = str(member.get('completed_points', 0) or 0)
        row.cells[3].text = f"{member.get('utilization', 0)}%"
    
    doc.add_paragraph(f'Workload Distribution: {improvement_areas["team_utilization"]["workload_distribution"]}')
    
    doc.add_paragraph('Optimization Suggestions:')
    for suggestion in improvement_areas['team_utilization']['optimization_suggestions']:
        doc.add_paragraph(suggestion, style='List Bullet')
    
    # Add Additional Improvements
    doc.add_heading('Additional Improvements', level=2)
    for improvement in improvement_areas['additional_improvements']:
        p = doc.add_paragraph()
        p.add_run(f'Area: {improvement["area"]}\n').bold = True
        p.add_run(f'Observation: {improvement["observation"]}\n')
        p.add_run(f'Suggestion: {improvement["suggestion"]}')
    
    return doc

@app.route('/api/sprint-combined-report', methods=['POST'])
def generate_combined_report():
    try:
        print("Starting combined report generation...")
        
        if 'file' not in request.files:
            print("No file in request")
            return jsonify({'error': 'No file provided'}), 400
        
        board_id = request.form.get('boardId')
        sprint_id = request.form.get('sprintId')
        
        print(f"Received board_id: {board_id}, sprint_id: {sprint_id}")
        
        if not board_id or not sprint_id:
            print("Missing board_id or sprint_id")
            return jsonify({'error': 'Board ID and Sprint ID are required'}), 400
        
        excel_file = request.files['file']
        if not excel_file.filename.endswith(('.xlsx', '.xls')):
            print(f"Invalid file format: {excel_file.filename}")
            return jsonify({'error': 'Invalid file format. Please upload an Excel file.'}), 400
        
        print("Getting Jira client...")
        # Get sprint details
        jira_client = get_jira_client()
        sprint = jira_client.sprint(sprint_id)
        if not sprint:
            print(f"Sprint not found: {sprint_id}")
            return jsonify({'error': 'Sprint not found'}), 404
        
        print("Getting sprint stories...")
        # Get sprint stories with all details
        sprint_stories = get_sprint_stories(jira_client, sprint_id)
        
        print("Processing Excel data...")
        # Process Excel data
        structured_data = process_excel_data(excel_file)
        
        print("Generating sprint data...")
        # Generate sprint data
        sprint_data = {
            'sprint_name': sprint.name,
            'sprint_goal': sprint.goal if hasattr(sprint, 'goal') else None,
            'start_date': sprint.startDate,
            'end_date': sprint.endDate,
            'stories': sprint_stories
        }
        
        print("Generating subgoals...")
        # Generate subgoals and achievements
        subgoals = generate_subgoals(sprint_data['sprint_goal'])
        
        print("Assigning stories to subgoals...")
        story_assignments = assign_stories_to_subgoals(sprint_stories, subgoals)
        
        print("Generating achievements...")
        achievements = generate_achievements(sprint_stories, subgoals)
        
        print("Generating improvement areas...")
        # Generate improvement areas
        improvement_areas = generate_improvement_areas(structured_data, sprint_data)
        
        print("Generating combined document...")
        try:
            # Generate combined document
            doc = generate_combined_sprint_doc(
                sprint_data,
                improvement_areas,
                subgoals,
                story_assignments,
                achievements,
                structured_data
            )
        except Exception as doc_error:
            print(f"Error in generate_combined_sprint_doc: {str(doc_error)}")
            raise
        
        print("Saving document...")
        # Save to BytesIO
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        print("Sending file...")
        return send_file(
            doc_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'sprint_report_and_analysis_{sprint.name}.docx'
        )
    
    except Exception as e:
        import traceback
        print(f"Error generating combined report: {str(e)}")
        print("Full traceback:")
        print(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True) 